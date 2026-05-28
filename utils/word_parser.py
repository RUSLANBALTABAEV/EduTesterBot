# utils/word_parser.py
"""
Парсер для извлечения тестов из Word документов.
Поддерживает формат:
1. Текст вопроса
A) Вариант ответа 1
B) Вариант ответа 2
C) Вариант ответа 3
D) Вариант ответа 4
"""

import re
import io
from typing import List, Dict, Tuple, Union
from docx import Document


class WordTestParser:
    """Парсер тестов из Word документа."""
    
    def __init__(self, source: Union[str, io.BytesIO]):
        """
        Инициализация парсера.
        
        Args:
            source: Путь к Word документу или BytesIO объект
        """
        if isinstance(source, str):
            self.doc = Document(source)
        elif isinstance(source, io.BytesIO):
            source.seek(0)
            self.doc = Document(source)
        else:
            raise ValueError("source должен быть строкой или BytesIO")
        
        self.tests = []
    
    def parse(self) -> List[Dict]:
        """
        Парсить весь документ и вернуть список тестов.
        
        Returns:
            Список тестов с вопросами
        """
        all_text = self._extract_paragraphs()
        self.questions = self._parse_questions(all_text)
        return self.questions
    
    def _extract_paragraphs(self) -> str:
        """Извлечь весь текст из документа."""
        text = ""
        for para in self.doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        return text
    
    def _parse_questions(self, text: str) -> List[Dict]:
        """
        Парсить вопросы из текста.
        
        Формат:
        N. Текст вопроса
        A) Вариант 1
        B) Вариант 2
        C) Вариант 3
        D) Вариант 4
        """
        questions = []
        lines = text.strip().split("\n")
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Ищем начало вопроса (N. Текст)
            if self._is_question_start(line):
                question = self._extract_question(lines, i)
                if question:
                    questions.append(question)
                    i += len(question.get('_lines_used', 1)) - 1
            
            i += 1
        
        return questions
    
    def _is_question_start(self, line: str) -> bool:
        """Проверить, начинается ли строка с номера вопроса."""
        pattern = r'^\d+\.\s+'
        return bool(re.match(pattern, line))
    
    def _extract_question(self, lines: List[str], start_idx: int) -> Dict | None:
        """
        Извлечь вопрос с вариантами ответа.
        
        Args:
            lines: Список строк документа
            start_idx: Индекс начала вопроса
            
        Returns:
            Словарь с вопросом и вариантами
        """
        question_line = lines[start_idx].strip()
        
        # Извлекаем номер и текст вопроса
        match = re.match(r'^(\d+)\.\s+(.+)$', question_line)
        if not match:
            return None
        
        question_num = int(match.group(1))
        question_text = match.group(2)
        
        options = []
        i = start_idx + 1
        correct_index = None
        
        # Ищем варианты ответов
        while i < len(lines):
            line = lines[i].strip()
            
            # Проверяем, является ли это вариантом ответа
            option_match = re.match(r'^([A-D])\)\s+(.+)$', line)
            
            if option_match:
                option_letter = option_match.group(1)
                option_text = option_match.group(2)
                
                # Определяем индекс (A=0, B=1, C=2, D=3)
                option_index = ord(option_letter) - ord('A')
                
                options.append({
                    'letter': option_letter,
                    'text': option_text,
                    'index': option_index
                })
                
                i += 1
            elif self._is_question_start(line) or line == "":
                # Конец текущего вопроса
                break
            else:
                i += 1
        
        if not options:
            return None
        
        # Первый вариант считаем правильным (как в примере документа)
        correct_index = 0
        
        result = {
            'number': question_num,
            'text': question_text,
            'type': 'single',  # По умолчанию один правильный ответ
            'points': 1.0,
            'options': options,
            'correct_index': correct_index,
            '_lines_used': i - start_idx
        }
        
        return result
    
    def get_questions_as_db_format(self) -> List[Dict]:
        """
        Получить вопросы в формате для сохранения в БД.
        
        Returns:
            Список вопросов в формате БД
        """
        db_questions = []
        
        for q in self.questions:
            db_question = {
                'text': q['text'],
                'type': q['type'],
                'points': q['points'],
                'order_num': q['number'],
                'options': []
            }
            
            # Добавляем варианты ответов
            for opt in q['options']:
                db_option = {
                    'text': opt['text'],
                    'is_correct': (opt['index'] == q['correct_index'])
                }
                db_question['options'].append(db_option)
            
            db_questions.append(db_question)
        
        return db_questions


def parse_word_file(filepath: str) -> Tuple[List[Dict], str]:
    """
    Парсить Word файл и вернуть список вопросов.
    
    Args:
        filepath: Путь к Word файлу
        
    Returns:
        Кортеж (список вопросов, сообщение об ошибке)
    """
    try:
        parser = WordTestParser(filepath)
        questions = parser.parse()
        
        if not questions:
            return [], "В документе не найдены вопросы в требуемом формате"
        
        db_format = parser.get_questions_as_db_format()
        return db_format, None
        
    except Exception as e:
        return [], f"Ошибка при парсинге файла: {str(e)}"
